from io import BytesIO
from docx import Document
import re


SECTION_NAMES = ["CONTACT", "SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION"]


def parse_match_score(analysis_text: str) -> int:
    if not analysis_text:
        return 0

    patterns = [
        r"Match Score:\s*-?\s*(\d{1,3})",
        r"Match Score\s*\n\s*-\s*(\d{1,3})",
        r"(\d{1,3})\s*/\s*100",
    ]

    for pattern in patterns:
        match = re.search(pattern, analysis_text, re.IGNORECASE)
        if match:
            score = int(match.group(1))
            return max(0, min(score, 100))

    return 0


def extract_match_summary(analysis_text: str) -> str:
    if not analysis_text:
        return "No analysis yet."

    lines = [line.strip() for line in analysis_text.splitlines() if line.strip()]
    capture = False
    collected = []

    for line in lines:
        if line.lower().startswith("match score"):
            capture = True
            continue

        if capture:
            if any(
                line.lower().startswith(section.lower())
                for section in [
                    "strengths",
                    "missing skills",
                    "improvement suggestions",
                    "rewritten resume bullets",
                ]
            ):
                break
            collected.append(line)

    if collected:
        return " ".join(collected)[:240]

    return "Resume-to-job match summary will appear here after analysis."


def score_label(score: int) -> str:
    if score >= 80:
        return "Strong Match"
    if score >= 60:
        return "Moderate Match"
    if score > 0:
        return "Needs Improvement"
    return "Not Available"


def parse_generated_resume_sections(text: str) -> dict:
    sections = {
        "name": "",
        "headline": "",
        "contact": "",
        "summary": "",
        "skills": "",
        "experience": "",
        "projects": "",
        "education": "",
    }

    if not text.strip():
        return sections

    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = [line.strip() for line in lines if line.strip()]

    if not cleaned:
        return sections

    sections["name"] = cleaned[0]

    idx = 1
    if idx < len(cleaned) and cleaned[idx].upper() not in SECTION_NAMES:
        sections["headline"] = cleaned[idx]
        idx += 1

    current_section = None

    for line in cleaned[idx:]:
        upper_line = line.upper()

        if upper_line in SECTION_NAMES:
            current_section = upper_line.lower()
            continue

        if current_section in sections:
            if sections[current_section]:
                sections[current_section] += "\n" + line
            else:
                sections[current_section] = line

    return sections


def build_resume_docx_from_sections(sections: dict) -> BytesIO:
    document = Document()

    name = sections.get("name", "").strip()
    headline = sections.get("headline", "").strip()

    if name:
        document.add_heading(name, level=0)
    if headline:
        document.add_paragraph(headline)

    def add_section(title: str, content: str):
        content = content.strip()
        if not content:
            return

        document.add_heading(title, level=1)
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

    add_section("Contact", sections.get("contact", ""))
    add_section("Summary", sections.get("summary", ""))
    add_section("Skills", sections.get("skills", ""))
    add_section("Experience", sections.get("experience", ""))
    add_section("Projects", sections.get("projects", ""))
    add_section("Education", sections.get("education", ""))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output